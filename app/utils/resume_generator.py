import os
from datetime import datetime
from docx import Document


def save_resume_docx(data: dict, template: str = "modern") -> str:
    """Generate and save resume as DOCX"""
    doc = Document()
    
    
    title = doc.add_paragraph()
    title_run = title.add_run(data.get('name', 'Resume'))
    title_run.bold = True
    title_run.font.size = 28000 
    title.alignment = 1  
    
    
    contact = doc.add_paragraph(
        f"{data.get('email', '')} | {data.get('phone', '')} | {data.get('linkedin', '')}"
    )
    contact.alignment = 1  # Center
    
    
    if data.get('summary'):
        doc.add_paragraph("PROFESSIONAL SUMMARY", style='Heading 1')
        doc.add_paragraph(data['summary'])
    
    
    if data.get('experience'):
        doc.add_paragraph("EXPERIENCE", style='Heading 1')
        for exp in data['experience']:
            p = doc.add_paragraph(f"{exp.get('title', '')} at {exp.get('company', '')}", style='Heading 2')
            doc.add_paragraph(f"Duration: {exp.get('duration', '')}")
            doc.add_paragraph(exp.get('description', ''))
    
    
    if data.get('education'):
        doc.add_paragraph("EDUCATION", style='Heading 1')
        for edu in data['education']:
            p = doc.add_paragraph(
                f"{edu.get('degree', '')} in {edu.get('field', '')}",
                style='Heading 2'
            )
            doc.add_paragraph(edu.get('institution', ''))
            doc.add_paragraph(f"Graduated: {edu.get('year', '')}")
    
   
    if data.get('skills'):
        doc.add_paragraph("SKILLS", style='Heading 1')
        skills_text = ", ".join(data['skills'])
        doc.add_paragraph(skills_text)
    
    
    if data.get('projects'):
        doc.add_paragraph("PROJECTS", style='Heading 1')
        for proj in data['projects']:
            doc.add_paragraph(proj.get('title', ''), style='Heading 2')
            doc.add_paragraph(proj.get('description', ''))
    
  
    os.makedirs('outputs', exist_ok=True)
    filename = f"resume_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    filepath = f"outputs/{filename}"
    doc.save(filepath)
    return filepath
