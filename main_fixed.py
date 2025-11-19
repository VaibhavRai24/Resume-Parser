from fastapi import FastAPI, File, UploadFile, Form
from fastapi.responses import FileResponse, JSONResponse
from fastapi.staticfiles import StaticFiles
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import Optional, List
import json
import os
from datetime import datetime
import PyPDF2
from docx import Document
from dotenv import load_dotenv
load_dotenv()
from langchain_groq import ChatGroq

from langchain.prompts import PromptTemplate
from langchain.chains import LLMChain
from langchain_core.output_parsers import JsonOutputParser
from langchain_core.exceptions import OutputParserException
from langgraph.graph import StateGraph
from typing_extensions import TypedDict

app = FastAPI(title="AI Resume Builder", version="1.0")


app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

app.mount("/static", StaticFiles(directory="static"), name="static")

# ==================== DATA MODELS ====================

class ResumeData(BaseModel):
    name: str
    email: str
    phone: str
    summary: str
    experience: List[dict]
    education: List[dict]
    skills: List[str]
    projects: Optional[List[dict]] = []

class ResumeState(TypedDict):
    """State for LangGraph workflow"""
    raw_text: str
    parsed_data: dict
    ats_score: dict
    enhanced_data: dict
    template: str
    output_file: str

# ==================== UTILITY FUNCTIONS ====================

def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from PDF file"""
    text = ""
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text += page.extract_text()
    except Exception as e:
        raise Exception(f"Error extracting PDF: {str(e)}")
    return text

def extract_text_from_docx(file_path: str) -> str:
    """Extract text from DOCX file"""
    try:
        doc = Document(file_path)
        text = "\n".join([para.text for para in doc.paragraphs])
    except Exception as e:
        raise Exception(f"Error extracting DOCX: {str(e)}")
    return text

def save_resume_docx(data: dict, template: str = "modern") -> str:
    """Generate and save resume as DOCX"""
    doc = Document()
    

    title = doc.add_paragraph()
    title_run = title.add_run(data.get('name', 'Resume'))
    title_run.bold = True
    title_run.font.size = 28000 
    title.alignment = 1  
    
    # Contact Info
    contact = doc.add_paragraph(
        f"{data.get('email', '')} | {data.get('phone', '')} | {data.get('linkedin', '')}"
    )
    contact.alignment = 1  # Center
    
  
    if data.get('summary'):
        doc.add_paragraph("PROFESSIONAL SUMMARY", style='Heading 1')
        doc.add_paragraph(data['summary'])
    
    # Experience
    if data.get('experience'):
        doc.add_paragraph("EXPERIENCE", style='Heading 1')
        for exp in data['experience']:
            p = doc.add_paragraph(f"{exp.get('title', '')} at {exp.get('company', '')}", style='Heading 2')
            doc.add_paragraph(f"Duration: {exp.get('duration', '')}")
            doc.add_paragraph(exp.get('description', ''))
    
    # Education
    if data.get('education'):
        doc.add_paragraph("EDUCATION", style='Heading 1')
        for edu in data['education']:
            p = doc.add_paragraph(
                f"{edu.get('degree', '')} in {edu.get('field', '')}",
                style='Heading 2'
            )
            doc.add_paragraph(edu.get('institution', ''))
            doc.add_paragraph(f"Graduated: {edu.get('year', '')}")
    
    # Skills
    if data.get('skills'):
        doc.add_paragraph("SKILLS", style='Heading 1')
        skills_text = ", ".join(data['skills'])
        doc.add_paragraph(skills_text)
    
    # Projects
    if data.get('projects'):
        doc.add_paragraph("PROJECTS", style='Heading 1')
        for proj in data['projects']:
            doc.add_paragraph(proj.get('title', ''), style='Heading 2')
            doc.add_paragraph(proj.get('description', ''))
    
   
    os.makedirs('outputs', exist_ok=True)
    filename = f"resume_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    filepath = f"outputs/{filename}"
    doc.save(filepath)
    return filepath

# ==================== LANGGRAPH WORKFLOW ====================

# Initialize LLM

llm = ChatGroq(model="llama-3.1-8b-instant", temperature=0.2)

# -------- Node: Parse Resume --------
def parse_resume_node(state: ResumeState) -> ResumeState:
    """Parse raw resume text into structured data"""
    
    # Add validation for empty text
    if not state["raw_text"] or len(state["raw_text"].strip()) < 20:
        raise Exception("Resume text is too short or empty to process")
    
    prompt = PromptTemplate(
        input_variables=["resume_text"],
        template="""
        Parse this resume and extract structured data. Return ONLY valid JSON.
        {resume_text}
        
        Return JSON with this structure:
        {{
            "name": "string",
            "email": "string",
            "phone": "string",
            "linkedin": "string",
            "summary": "string",
            "experience": [
                {{"title": "string", "company": "string", "duration": "string", "description": "string"}}
            ],
            "education": [
                {{"degree": "string", "field": "string", "institution": "string", "year": "string"}}
            ],
            "skills": ["string"],
            "projects": [
                {{"title": "string", "description": "string"}}
            ]
        }}
        """
    )
    
    parser = JsonOutputParser()
    chain = prompt | llm | parser
    
    try:
        parsed_data = chain.invoke({"resume_text": state["raw_text"]})
    except OutputParserException as e:
        raise Exception(f"Failed to parse resume as JSON: {str(e)}")
    
    state["parsed_data"] = parsed_data
    return state

# -------- Node: Calculate ATS Score --------
def ats_score_node(state: ResumeState) -> ResumeState:
    """Calculate ATS score for resume"""
    prompt = PromptTemplate(
        input_variables=["resume_data"],
        template="""
        Analyze this resume data for ATS (Applicant Tracking System) compatibility.
        {resume_data}
        
        Return ONLY valid JSON with:
        {{
            "score": <0-100>,
            "feedback": "string",
            "improvements": ["string list of improvements"],
            "missing_keywords": ["string list of keywords to add"]
        }}
        """
    )
    
    parser = JsonOutputParser()
    chain = prompt | llm | parser
    
    try:
        ats_data = chain.invoke({"resume_data": json.dumps(state["parsed_data"])})
    except OutputParserException as e:
        ats_data = {
            "score": 0,
            "feedback": "Could not calculate ATS score",
            "improvements": [],
            "missing_keywords": []
        }
    
    state["ats_score"] = ats_data
    return state

# -------- Node: Enhance Resume --------
def enhance_resume_node(state: ResumeState) -> ResumeState:
    """Enhance resume with AI improvements"""
    prompt = PromptTemplate(
        input_variables=["resume_data", "ats_feedback"],
        template="""
        Improve this resume to increase ATS score and professional impact.
        Resume: {resume_data}
        ATS Feedback: {ats_feedback}
        
        Return ONLY valid JSON with enhanced resume structure:
        {{
            "name": "string",
            "email": "string",
            "phone": "string",
            "linkedin": "string",
            "summary": "improved professional summary",
            "experience": [
                {{"title": "string", "company": "string", "duration": "string", "description": "enhanced description with action verbs"}}
            ],
            "education": [
                {{"degree": "string", "field": "string", "institution": "string", "year": "string"}}
            ],
            "skills": ["enhanced skill list"],
            "projects": [
                {{"title": "string", "description": "enhanced description"}}
            ]
        }}
        """
    )
    
    parser = JsonOutputParser()
    chain = prompt | llm | parser
    
    try:
        enhanced_data = chain.invoke({
            "resume_data": json.dumps(state["parsed_data"]),
            "ats_feedback": json.dumps(state["ats_score"])
        })
    except OutputParserException as e:
        enhanced_data = state["parsed_data"]
    
    state["enhanced_data"] = enhanced_data
    return state

# -------- Node: Generate Resume --------
def generate_resume_node(state: ResumeState) -> ResumeState:
    """Generate final resume file"""
   
    resume_data = state.get("enhanced_data") or state.get("parsed_data")
    template = state.get("template", "modern")
    
    filepath = save_resume_docx(resume_data, template)
    state["output_file"] = filepath
    return state

# -------- Build LangGraph --------
workflow = StateGraph(ResumeState)

# Add nodes
workflow.add_node("parse", parse_resume_node)
workflow.add_node("ats_score_analysis", ats_score_node)
workflow.add_node("enhance", enhance_resume_node)
workflow.add_node("generate", generate_resume_node)

workflow.add_edge("parse", "ats_score_analysis")
workflow.add_edge("ats_score_analysis", "enhance")
workflow.add_edge("enhance", "generate")

workflow.set_entry_point("parse")
workflow.set_finish_point("generate")

# Compile
graph = workflow.compile()

# ==================== API ENDPOINTS ====================

@app.get("/")
async def root():
    """Serve main HTML page"""
    return FileResponse("static/index.html")

@app.post("/api/upload")
async def upload_resume(file: UploadFile = File(...)):
    """Upload and process resume file"""
    try:

        os.makedirs("uploads", exist_ok=True)
        filepath = f"uploads/{file.filename}"
        with open(filepath, "wb") as f:
            content = await file.read()
            f.write(content)
        
       
        if file.filename.endswith(".pdf"):
            raw_text = extract_text_from_pdf(filepath)
        elif file.filename.endswith(".docx"):
            raw_text = extract_text_from_docx(filepath)
        else:
            return JSONResponse({"error": "Unsupported file format"}, status_code=400)
        
        # Run LangGraph workflow
        initial_state = {
            "raw_text": raw_text,
            "parsed_data": {},
            "ats_score": {},
            "enhanced_data": {},
            "template": "modern",
            "output_file": ""
        }
        
        result = graph.invoke(initial_state)
        
        # Extract just the filename from the full path
        output_filename = os.path.basename(result["output_file"])
        
        return JSONResponse({
            "status": "success",
            "parsed_data": result["parsed_data"],
            "ats_score": result["ats_score"],
            "output_file": output_filename
        })
    
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=400)

@app.post("/api/process-manual")
async def process_manual_resume(data: ResumeData):
    """Process manually entered resume data"""
    try:
        raw_text = f"""
        Name: {data.name}
        Email: {data.email}
        Phone: {data.phone}
        
        Summary: {data.summary}
        
        Experience: {json.dumps(data.experience)}
        Education: {json.dumps(data.education)}
        Skills: {', '.join(data.skills)}
        Projects: {json.dumps(data.projects)}
        """
        
        initial_state = {
            "raw_text": raw_text,
            "parsed_data": data.dict(),
            "ats_score": {},
            "enhanced_data": {},
            "template": "modern",
            "output_file": ""
        }
        
        result = graph.invoke(initial_state)
        
        output_filename = os.path.basename(result["output_file"])
        
        return JSONResponse({
            "status": "success",
            "parsed_data": result["parsed_data"],
            "ats_score": result["ats_score"],
            "output_file": output_filename
        })
    
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=400)

@app.post("/api/enhance")
async def enhance_resume(data: dict):
    """Enhance existing resume"""
    try:
        raw_text = json.dumps(data)
        
        initial_state = {
            "raw_text": raw_text,
            "parsed_data": data,
            "ats_score": {},
            "enhanced_data": {},
            "template": "modern",
            "output_file": ""
        }
        
        # Run from ats_score onwards
        result = graph.invoke(initial_state)
        
        return JSONResponse({
            "status": "success",
            "enhanced_data": result["enhanced_data"],
            "ats_score": result["ats_score"]
        })
    
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=400)

@app.get("/api/download/{filename}")
async def download_resume(filename: str):
    """Download generated resume"""
    try:
        filepath = f"outputs/{filename}"
        if os.path.exists(filepath):
            return FileResponse(filepath, filename=filename)
        return JSONResponse({"error": "File not found"}, status_code=404)
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=400)

@app.get("/api/health")
async def health():
    """Health check"""
    return {"status": "ok", "service": "AI Resume Builder"}

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
